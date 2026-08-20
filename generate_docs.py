from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='SimSun', font_size=10.5, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_chinese(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        set_chinese_font(run, 'SimHei', 16, True)
    elif level == 2:
        set_chinese_font(run, 'SimHei', 14, True)
    else:
        set_chinese_font(run, 'SimHei', 12, True)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return heading

def add_paragraph_chinese(doc, text, bold=False, size=10.5, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, 'SimSun', size, bold)
    p.alignment = alignment
    return p

# ==================== 生成技术路线规划书 ====================
doc1 = Document()

# 标题
title = doc1.add_heading(level=0)
title_run = title.add_run('电网卫士——面向特种工况的低能耗抗污染反渗透膜系统\n技术路线规划书')
set_chinese_font(title_run, 'SimHei', 18, True)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 一、项目概述
add_heading_chinese(doc1, '一、项目概述', 1)
add_paragraph_chinese(doc1, 
    '本项目针对特高压直流输电换流站阀冷系统反渗透膜长期依赖进口、高海拔差水质工况下污堵严重、'
    '运维模式被动等痛点，构建"基础研究-技术开发-产业应用"全链条创新体系，开发具有自主知识产权的'
    '电网专用反渗透膜及智能运维系统。')

add_paragraph_chinese(doc1, '核心技术目标：', bold=True)
goals = [
    '• 脱盐率≥99.7%，寿命≥5年',
    '• 能耗降低20%以上，清洗周期延长至6-8个月',
    '• 实现膜材料"单体-配方-工艺"全链条自主可控'
]
for goal in goals:
    add_paragraph_chinese(doc1, goal)

# 二、总体技术路线
add_heading_chinese(doc1, '二、总体技术路线', 1)
add_paragraph_chinese(doc1, 
    '采用"理论指导-技术突破-工程验证-产业转化"四阶段递进策略，预计周期4年（2024-2028年）。')

route_text = '''基础研究阶段（第1-2年）
    ↓ 理论输出
技术开发阶段（第2-3年）
    ↓ 产品雏形
中试验证阶段（第3-4年）
    ↓ 工程数据
产业化阶段（第4年起）'''
add_paragraph_chinese(doc1, route_text)

# 三、分阶段技术路线
add_heading_chinese(doc1, '三、分阶段技术路线', 1)

# 阶段一
add_heading_chinese(doc1, '3.1 基础研究阶段（2024.01-2025.12）', 2)
add_paragraph_chinese(doc1, '核心任务：揭示聚酰胺微观结构与性能的关联机制，建立分子设计理论指导体系', bold=True)

table1 = doc1.add_table(rows=5, cols=5)
table1.style = 'Light Grid Accent 1'
headers = ['序号', '研究内容', '技术路线', '预期成果', '时间节点']
for i, header in enumerate(headers):
    cell = table1.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 10, True)

data1 = [
    ['1.1', '聚酰胺微观结构表征', '采用分子动力学模拟（MD）结合PALS、SAXS等表征手段', '建立结构-性能关联模型1套', '2024.12'],
    ['1.2', '低温/高污染工况失效机理', '研究聚酰胺膜在5-15℃低温、高浊度进水条件下的失效规律', '形成失效机理研究报告', '2025.03'],
    ['1.3', '分子设计理论指导', '基于结构-性能关系，设计低温高活性胺类单体分子结构', '筛选优化单体分子2-3种', '2025.09'],
    ['1.4', '配方优化理论模型', '建立"单体配比-反应条件-膜性能"多因素耦合优化模型', '建立配方设计理论模型', '2025.12']
]
for i, row_data in enumerate(data1, 1):
    for j, cell_data in enumerate(row_data):
        cell = table1.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

add_paragraph_chinese(doc1, '阶段里程碑：完成聚酰胺膜结构-性能基础理论研究，形成分子设计理论指导体系，发表SCI/EI论文2-3篇，申请发明专利1-2项。')

# 阶段二
add_heading_chinese(doc1, '3.2 技术开发阶段（2025.06-2026.12）', 2)
add_paragraph_chinese(doc1, '核心任务：突破高纯单体合成、抗污染-耐低温协同改性、CFD流道优化等关键技术', bold=True)

table2 = doc1.add_table(rows=6, cols=5)
table2.style = 'Light Grid Accent 1'
for i, header in enumerate(headers):
    cell = table2.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 10, True)

data2 = [
    ['2.1', '高纯单体合成技术', '开发MPD及其衍生物的高纯度合成工艺（纯度≥99.5%）', '高纯单体合成工艺1套', '2025.12'],
    ['2.2', '界面聚合配方技术', '优化水相/油相单体浓度、添加剂配比、反应温度等参数', '确定最优配方组合1-2组', '2026.03'],
    ['2.3', '抗污染-耐低温协同改性', '两性离子聚合物表面接枝+TiO₂光催化纳米粒子共混+超亲水涂层', '改性膜样品，接触角<30°', '2026.06'],
    ['2.4', 'CFD流道优化设计', '采用ANSYS Fluent进行流场模拟，优化进水隔网几何结构', '压降降低20%', '2026.09'],
    ['2.5', '膜片实验室制备', '建立实验室制膜平台（涂布面积≥0.5m²）', '实验室膜片脱盐率≥99.5%', '2026.12']
]
for i, row_data in enumerate(data2, 1):
    for j, cell_data in enumerate(row_data):
        cell = table2.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

add_paragraph_chinese(doc1, '阶段里程碑：完成高纯单体合成工艺开发，突破抗污染-耐低温协同改性技术，形成CFD优化流道设计方案，申请发明专利3-5项。')

# 阶段三
add_heading_chinese(doc1, '3.3 中试验证阶段（2026.09-2027.12）', 2)
add_paragraph_chinese(doc1, '核心任务：在典型换流站开展现场中试验证，建立寿命预测模型，制定行业标准', bold=True)

table3 = doc1.add_table(rows=6, cols=5)
table3.style = 'Light Grid Accent 1'
for i, header in enumerate(headers):
    cell = table3.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 10, True)

data3 = [
    ['3.1', '千支级中试线建设', '设计建设卷式反渗透膜中试生产线（产能≥1000支/年）', '建成中试生产线1条', '2027.03'],
    ['3.2', '中试膜元件制备', '采用中试线制备4英寸/8英寸膜元件', '良品率≥85%', '2027.06'],
    ['3.3', '现场中试验证', '选择2-3个典型换流站开展6-12个月现场挂片试验', '清洗周期≥6个月', '2027.12'],
    ['3.4', '膜寿命预测模型', '采用LSTM/GRU深度学习算法构建膜性能衰减预测模型', '预测误差<15%', '2027.09'],
    ['3.5', '行业标准制定', '牵头制定《电网用反渗透膜元件技术条件》等标准', '形成标准草案2-3项', '2027.12']
]
for i, row_data in enumerate(data3, 1):
    for j, cell_data in enumerate(row_data):
        cell = table3.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

add_paragraph_chinese(doc1, '阶段里程碑：建成千支级中试线并完成现场验证，膜性能达到预期目标（脱盐率≥99.7%、寿命≥5年），建立膜寿命预测模型。')

# 阶段四
add_heading_chinese(doc1, '3.4 产业化阶段（2027.06-2028.12及以后）', 2)
add_paragraph_chinese(doc1, '核心任务：实现万支级规模化量产，建立市场推广体系，形成完整产业链', bold=True)

table4 = doc1.add_table(rows=5, cols=5)
table4.style = 'Light Grid Accent 1'
for i, header in enumerate(headers):
    cell = table4.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 10, True)

data4 = [
    ['4.1', '万支级量产工艺', '设计建设万支级生产线（产能≥10000支/年）', '建成万支级生产线', '2028.06'],
    ['4.2', '质量管控体系', '建立全流程质量管理体系', '通过ISO9001、ISO14001认证', '2028.03'],
    ['4.3', '智能运维系统', '开发基于寿命预测模型的膜系统智能运维平台', '智能运维系统软件1套', '2028.09'],
    ['4.4', '市场推广与应用', '与国网、南网建立合作关系，开展首批次工程应用', '首批次工程应用≥5个换流站', '2028.12']
]
for i, row_data in enumerate(data4, 1):
    for j, cell_data in enumerate(row_data):
        cell = table4.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

# 四、里程碑节点
add_heading_chinese(doc1, '四、关键技术节点与里程碑', 1)

table5 = doc1.add_table(rows=9, cols=3)
table5.style = 'Light Grid Accent 1'
milestone_headers = ['时间节点', '里程碑事件', '标志性成果']
for i, header in enumerate(milestone_headers):
    cell = table5.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 10, True)

milestones = [
    ['2024.12', '基础理论突破', '建立结构-性能关联模型'],
    ['2025.12', '单体合成突破', '高纯单体合成工艺定型'],
    ['2026.06', '改性技术突破', '抗污染-耐低温改性膜样品'],
    ['2026.12', '实验室样品', '实验室膜片脱盐率≥99.5%'],
    ['2027.06', '中试线建成', '千支级中试线投产'],
    ['2027.12', '现场验证完成', '现场试验报告、行业标准草案'],
    ['2028.06', '量产线建成', '万支级生产线投产'],
    ['2028.12', '首单落地', '首批次工程应用合同']
]
for i, row_data in enumerate(milestones, 1):
    for j, cell_data in enumerate(row_data):
        cell = table5.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

# 五、技术风险与应对措施
add_heading_chinese(doc1, '五、技术风险与应对措施', 1)

table6 = doc1.add_table(rows=5, cols=3)
table6.style = 'Light Grid Accent 1'
risk_headers = ['风险类型', '风险描述', '应对措施']
for i, header in enumerate(risk_headers):
    cell = table6.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 10, True)

risks = [
    ['技术风险', '低温工况下聚酰胺膜通量衰减超预期', '①备选单体分子设计\n②引入亲水性中间层\n③优化支撑层结构'],
    ['工艺风险', '中试放大过程性能一致性难以保证', '①DOE实验设计优化\n②建立SPC统计控制\n③关键设备定制化'],
    ['验证风险', '现场工况复杂，试验周期延长', '①多站点并行试验\n②加速老化实验补充\n③联合攻关机制'],
    ['市场风险', '电网采购周期长，推广进度不及预期', '①提前介入标准制定\n②建立示范工程\n③全生命周期成本论证']
]
for i, row_data in enumerate(risks, 1):
    for j, cell_data in enumerate(row_data):
        cell = table6.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

# 六、知识产权规划
add_heading_chinese(doc1, '六、知识产权规划', 1)
add_paragraph_chinese(doc1, '采用"核心专利+外围专利"组合布局策略：', bold=True)
add_paragraph_chinese(doc1, '核心专利：高纯单体合成工艺、抗污染-耐低温协同改性方法、CFD优化流道结构')
add_paragraph_chinese(doc1, '外围专利：具体配方组合、工艺参数范围、专用设备、预测算法、应用场景')

# 页脚
add_paragraph_chinese(doc1, '\n\n')
add_paragraph_chinese(doc1, '编制单位：[参赛单位名称]', size=9)
add_paragraph_chinese(doc1, '编制日期：2026年3月', size=9)
add_paragraph_chinese(doc1, '版本号：V1.0', size=9)
add_paragraph_chinese(doc1, '注：本规划书基于创意阶段技术设想编制，具体参数以实际研发结果为准。', size=9)

# 保存第一个文档
doc1.save('/root/.openclaw/workspace/技术路线规划书.docx')
print("✓ 技术路线规划书.docx 已生成")

# ==================== 生成数值模拟方案 ====================
doc2 = Document()

# 标题
title2 = doc2.add_heading(level=0)
title_run2 = title2.add_run('电网专用反渗透膜CFD流道优化\n数值模拟方案')
set_chinese_font(title_run2, 'SimHei', 18, True)
title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 一、方案概述
add_heading_chinese(doc2, '一、方案概述', 1)
add_paragraph_chinese(doc2, 
    '本方案针对换流站阀冷系统反渗透膜组件流道结构开展计算流体力学（CFD）数值模拟研究，'
    '通过优化进水隔网几何参数，降低膜组件运行压降与浓差极化，预期实现能耗降低20%以上的技术目标。')

add_paragraph_chinese(doc2, '模拟目标：', bold=True)
targets = [
    '• 降低膜组件运行压降≥20%',
    '• 减少浓差极化系数（CP）≥15%',
    '• 提高流道内流速均匀性（不均匀系数<0.3）',
    '• 增强流道自清洁能力，减少死区'
]
for target in targets:
    add_paragraph_chinese(doc2, target)

# 二、模拟对象与工况
add_heading_chinese(doc2, '二、模拟对象与工况', 1)

add_heading_chinese(doc2, '2.1 模拟对象', 2)
add_paragraph_chinese(doc2, '研究对象：4040型（4英寸）卷式反渗透膜组件进水隔网')
add_paragraph_chinese(doc2, '隔网类型：')
add_paragraph_chinese(doc2, '• 传统菱形隔网（Commercial spacer，基准方案）')
add_paragraph_chinese(doc2, '• 优化波纹隔网（Wave-type spacer，仿生设计）')
add_paragraph_chinese(doc2, '• 交错网格隔网（Intersecting grid，低污染设计）')

add_heading_chinese(doc2, '2.2 设计参数变量', 2)

table7 = doc2.add_table(rows=6, cols=4)
table7.style = 'Light Grid Accent 1'
param_headers = ['参数名称', '符号', '取值范围', '单位']
for i, header in enumerate(param_headers):
    cell = table7.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 10, True)

params = [
    ['隔网厚度', 'h', '0.6 - 1.0', 'mm'],
    ['菱形长轴', 'L1', '3.0 - 5.0', 'mm'],
    ['菱形短轴', 'L2', '1.5 - 3.0', 'mm'],
    ['隔网角度', 'θ', '45 - 90', '°'],
    ['波纹振幅', 'A', '0.2 - 0.5', 'mm']
]
for i, row_data in enumerate(params, 1):
    for j, cell_data in enumerate(row_data):
        cell = table7.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

add_heading_chinese(doc2, '2.3 运行工况', 2)
conditions = [
    '• 进水流量：Q = 0.5 - 2.0 m³/h（典型工况：1.0 m³/h）',
    '• 操作压力：P = 0.8 - 1.5 MPa',
    '• 进水温度：T = 5 - 25℃（重点研究低温工况5-10℃）',
    '• 进水含盐量：C = 500 - 2000 mg/L（以NaCl计）',
    '• 回收率：Y = 15 - 25%'
]
for cond in conditions:
    add_paragraph_chinese(doc2, cond)

# 三、数学模型
add_heading_chinese(doc2, '三、数学模型与数值方法', 1)

add_heading_chinese(doc2, '3.1 控制方程', 2)
add_paragraph_chinese(doc2, '采用稳态不可压缩Navier-Stokes方程描述流场：', bold=True)
add_paragraph_chinese(doc2, '连续性方程：∇·(ρu) = 0')
add_paragraph_chinese(doc2, '动量方程：∇·(ρuu) = -∇p + ∇·[μ(∇u + (∇u)ᵀ)] + Sₘ')
add_paragraph_chinese(doc2, '其中：ρ为流体密度，u为速度矢量，p为压力，μ为动力粘度，Sₘ为源项')

add_paragraph_chinese(doc2, '溶质输运方程（考虑浓差极化）：', bold=True)
add_paragraph_chinese(doc2, '∇·(uC) = ∇·(D_eff∇C) + S_c')
add_paragraph_chinese(doc2, '其中：C为溶质浓度，D_eff为有效扩散系数，S_c为溶质源项')

add_heading_chinese(doc2, '3.2 边界条件', 2)
table8 = doc2.add_table(rows=5, cols=3)
table8.style = 'Light Grid Accent 1'
bc_headers = ['边界位置', '边界类型', '具体设置']
for i, header in enumerate(bc_headers):
    cell = table8.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 10, True)

bcs = [
    ['进水口', '速度入口', '均匀流速 u = Q/A_in，浓度 C = C_in'],
    ['出水口', '压力出口', '表压 P = 0 Pa，充分发展流动'],
    ['膜表面', '壁面', '无滑移条件，通量 J = A(ΔP - Δπ)，溶质截留率 R'],
    ['隔网表面', '壁面', '无滑移条件，无渗透']
]
for i, row_data in enumerate(bcs, 1):
    for j, cell_data in enumerate(row_data):
        cell = table8.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

add_heading_chinese(doc2, '3.3 数值求解方法', 2)
methods = [
    '• 求解器：ANSYS Fluent 2023 R1',
    '• 湍流模型：标准k-ε模型（高雷诺数）/ 层流模型（低雷诺数验证）',
    '• 离散格式：二阶迎风格式（对流项），二阶中心差分（扩散项）',
    '• 压力-速度耦合：SIMPLE算法',
    '• 收敛判据：残差<10⁻⁶，进出口质量流量偏差<0.1%'
]
for method in methods:
    add_paragraph_chinese(doc2, method)

# 四、网格划分策略
add_heading_chinese(doc2, '四、网格划分策略', 1)

add_heading_chinese(doc2, '4.1 几何建模', 2)
add_paragraph_chinese(doc2, '采用周期性单元法简化计算域：选取单个菱形隔网单元作为代表性体积元（RVE），'
                            '设置周期性边界条件模拟无限延伸的隔网结构。计算域尺寸：L×W×H = 5mm × 3mm × 1mm')

add_heading_chinese(doc2, '4.2 网格类型与参数', 2)
table9 = doc2.add_table(rows=4, cols=4)
table9.style = 'Light Grid Accent 1'
grid_headers = ['区域', '网格类型', '网格尺寸', '备注']
for i, header in enumerate(grid_headers):
    cell = table9.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 10, True)

grids = [
    ['流道主体', '非结构化四面体', '0.05 - 0.15 mm', '局部加密'],
    ['隔网表面', '边界层网格', '首层高度0.01mm，增长比1.2，5层', 'y⁺<1'],
    ['膜表面', '边界层网格', '首层高度0.005mm，增长比1.1，8层', '解析浓差极化层']
]
for i, row_data in enumerate(grids, 1):
    for j, cell_data in enumerate(row_data):
        cell = table9.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

add_heading_chinese(doc2, '4.3 网格无关性验证', 2)
add_paragraph_chinese(doc2, '采用三套网格（粗：50万，中：120万，细：280万）进行无关性验证，'
                            '以压降和膜表面平均浓度为考核指标，确定最终计算网格（预计120-150万单元）。')

# 五、评价指标
add_heading_chinese(doc2, '五、评价指标与优化目标', 1)

table10 = doc2.add_table(rows=6, cols=3)
table10.style = 'Light Grid Accent 1'
metric_headers = ['评价指标', '计算公式', '优化目标']
for i, header in enumerate(metric_headers):
    cell = table10.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 10, True)

metrics = [
    ['压降（ΔP）', 'ΔP = P_in - P_out', '降低≥20%'],
    ['浓差极化系数（CP）', 'CP = C_wall/C_bulk', '降低≥15%'],
    ['流速不均匀系数（γ）', 'γ = σ_u / ū', '<0.3'],
    ['剪切速率（τ）', 'τ = μ(du/dy)', '提高以增强自清洁'],
    ['能耗指数（E）', 'E = ΔP·Q / (J·A)', '降低≥20%']
]
for i, row_data in enumerate(metrics, 1):
    for j, cell_data in enumerate(row_data):
        cell = table10.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

# 六、优化策略
add_heading_chinese(doc2, '六、优化策略与步骤', 1)

add_heading_chinese(doc2, '6.1 单因素敏感性分析', 2)
add_paragraph_chinese(doc2, '依次研究隔网厚度、角度、波纹振幅等单一参数对压降和浓差极化的影响规律，'
                            '确定各参数的敏感区间。')

add_heading_chinese(doc2, '6.2 多因素优化设计', 2)
add_paragraph_chinese(doc2, '采用响应面法（RSM）或遗传算法（GA）进行多参数耦合优化：')
add_paragraph_chinese(doc2, '• 设计变量：h, L1, L2, θ, A')
add_paragraph_chinese(doc2, '• 约束条件：结构可制造性、机械强度、压降上限')
add_paragraph_chinese(doc2, '• 目标函数：min f = w₁·ΔP/ΔP₀ + w₂·CP/CP₀（权重w₁=w₂=0.5）')

add_heading_chinese(doc2, '6.3 优化验证', 2)
add_paragraph_chinese(doc2, '对优化后的隔网结构进行详细流场分析，验证：')
add_paragraph_chinese(doc2, '• 速度分布云图：识别死区和高速区')
add_paragraph_chinese(doc2, '• 浓度分布云图：评估浓差极化改善效果')
add_paragraph_chinese(doc2, '• 湍动能分布：分析涡流强化混合效果')
add_paragraph_chinese(doc2, '• 粒子追踪：模拟污染物在流道中的运动轨迹')

# 七、预期成果
add_heading_chinese(doc2, '七、预期成果', 1)

outputs = [
    '1. 最优隔网结构参数组合1-2组（几何尺寸+性能指标）',
    '2. 流场特性分析报告（速度场、压力场、浓度场可视化结果）',
    '3. 与传统隔网的性能对比报告（压降、能耗、抗污染性）',
    '4. 隔网结构优化设计指导手册',
    '5. CFD模型文件及网格文件（可复现）'
]
for output in outputs:
    add_paragraph_chinese(doc2, output)

# 八、进度安排
add_heading_chinese(doc2, '八、进度安排', 1)

table11 = doc2.add_table(rows=6, cols=3)
table11.style = 'Light Grid Accent 1'
schedule_headers = ['阶段', '时间', '主要任务']
for i, header in enumerate(schedule_headers):
    cell = table11.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 10, True)

schedules = [
    ['几何建模与网格', '2026.01-2026.03', '三维建模、网格划分、无关性验证'],
    ['基准方案模拟', '2026.04-2026.05', '传统隔网流场模拟、基准数据建立'],
    ['单因素分析', '2026.06-2026.07', '敏感性分析、参数影响规律'],
    ['多因素优化', '2026.08-2026.09', '响应面优化、最优方案确定'],
    ['验证与总结', '2026.10-2026.12', '详细分析、报告撰写、模型归档']
]
for i, row_data in enumerate(schedules, 1):
    for j, cell_data in enumerate(row_data):
        cell = table11.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

# 九、资源需求
add_heading_chinese(doc2, '九、资源需求', 1)

resources = [
    '• 软件：ANSYS Fluent 2023 R1（或同等CFD软件）',
    '• 硬件：高性能计算工作站（CPU：32核以上，内存：128GB以上）',
    '• 人员：CFD仿真工程师1名，膜技术背景工程师1名',
    '• 周期：12个月'
]
for resource in resources:
    add_paragraph_chinese(doc2, resource)

# 页脚
add_paragraph_chinese(doc2, '\n\n')
add_paragraph_chinese(doc2, '编制单位：[参赛单位名称]', size=9)
add_paragraph_chinese(doc2, '编制日期：2026年3月', size=9)
add_paragraph_chinese(doc2, '版本号：V1.0', size=9)
add_paragraph_chinese(doc2, '注：本方案基于创意阶段技术设想编制，具体参数以实际仿真结果为准。', size=9)

# 保存第二个文档
doc2.save('/root/.openclaw/workspace/数值模拟方案.docx')
print("✓ 数值模拟方案.docx 已生成")

print("\n✅ 两个Word文档已全部生成！")
